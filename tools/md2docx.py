#!/usr/bin/env python
#
# file: tools/md2docx.py
#
# Convert a Markdown file to a Word .docx on the command line.
# No pandoc needed -- uses python-docx only.
#
# Handles: ATX headings (#..######), **bold**, *italic* / _italic_,
# `inline code`, [links](url), bullet and numbered lists, pipe tables,
# fenced ``` code blocks, > blockquotes, and --- horizontal rules.
#
# Usage:
#   python tools/md2docx.py input.md [output.docx]
#   python tools/md2docx.py paper/PREREGISTRATION.md
#------------------------------------------------------------------------------
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

# inline tokens, tried left-to-right at each position
#
_INLINE = re.compile(
    r"(?P<code>`[^`]+`)"
    r"|(?P<bold>\*\*[^*]+\*\*)"
    r"|(?P<italic>\*[^*]+\*|_[^_]+_)"
    r"|(?P<link>\[[^\]]+\]\([^)]+\))"
)


def add_runs(paragraph, text: str) -> None:
    """Add inline-formatted runs (bold/italic/code/link) to a paragraph."""
    pos = 0
    for m in _INLINE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        kind = m.lastgroup
        tok = m.group()
        if kind == "code":
            r = paragraph.add_run(tok[1:-1])
            r.font.name = "Consolas"
            r.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        elif kind == "bold":
            paragraph.add_run(tok[2:-2]).bold = True
        elif kind == "italic":
            paragraph.add_run(tok[1:-1]).italic = True
        elif kind == "link":
            lm = re.match(r"\[([^\]]+)\]\(([^)]+)\)", tok)
            label, url = lm.group(1), lm.group(2)
            r = paragraph.add_run(label)
            r.font.color.rgb = RGBColor(0x1A, 0x0D, 0xAB)
            r.underline = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_table(doc, rows: list) -> None:
    """rows: list of cell-lists; first row is the header."""
    cols = max(len(r) for r in rows)
    t = doc.add_table(rows=0, cols=cols)
    t.style = "Light Grid Accent 1"
    for i, cells in enumerate(rows):
        cells = cells + [""] * (cols - len(cells))
        tr = t.add_row().cells
        for j, cell in enumerate(cells):
            tr[j].paragraphs[0].text = ""
            add_runs(tr[j].paragraphs[0], cell.strip())
            if i == 0:
                for run in tr[j].paragraphs[0].runs:
                    run.bold = True


def split_row(line: str) -> list:
    """Split a pipe-table row into trimmed cells."""
    return [c.strip() for c in line.strip().strip("|").split("|")]


def is_sep(line: str) -> bool:
    """True for a table separator row like |---|:--:|."""
    return bool(re.match(r"^\s*\|?\s*:?-{3,}", line)) and "-" in line


def convert(md_path: Path, out_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    i, n = 0, len(lines)
    while i < n:
        line = lines[i]

        # fenced code block
        if line.lstrip().startswith("```"):
            i += 1
            buf = []
            while i < n and not lines[i].lstrip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1
            p = doc.add_paragraph()
            r = p.add_run("\n".join(buf))
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
            continue

        # heading
        h = re.match(r"^(#{1,6})\s+(.*)$", line)
        if h:
            doc.add_heading(h.group(2).strip(), level=min(len(h.group(1)), 4))
            i += 1
            continue

        # horizontal rule
        if re.match(r"^\s*([-*_])\1{2,}\s*$", line):
            doc.add_paragraph().add_run("_" * 40).font.color.rgb = \
                RGBColor(0xBB, 0xBB, 0xBB)
            i += 1
            continue

        # pipe table (header line + separator)
        if line.strip().startswith("|") and i + 1 < n and is_sep(lines[i + 1]):
            rows = [split_row(line)]
            i += 2
            while i < n and lines[i].strip().startswith("|"):
                rows.append(split_row(lines[i]))
                i += 1
            add_table(doc, rows)
            continue

        # blockquote
        if line.lstrip().startswith(">"):
            p = doc.add_paragraph(style="Intense Quote")
            add_runs(p, line.lstrip()[1:].strip())
            i += 1
            continue

        # bullet list
        b = re.match(r"^\s*[-*+]\s+(.*)$", line)
        if b:
            add_runs(doc.add_paragraph(style="List Bullet"), b.group(1))
            i += 1
            continue

        # numbered list
        num = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if num:
            add_runs(doc.add_paragraph(style="List Number"), num.group(1))
            i += 1
            continue

        # blank line
        if not line.strip():
            i += 1
            continue

        # paragraph (merge consecutive non-blank, non-special lines)
        buf = [line]
        i += 1
        while i < n and lines[i].strip() and not re.match(
            r"^(#{1,6}\s|\s*[-*+]\s|\s*\d+\.\s|>|\|)|^```", lines[i]
        ):
            buf.append(lines[i])
            i += 1
        add_runs(doc.add_paragraph(), " ".join(s.strip() for s in buf))

    doc.save(out_path)
    print(f"wrote {out_path}")


def main() -> int:
    if len(sys.argv) < 2:
        print("usage: python tools/md2docx.py input.md [output.docx]")
        return 2
    md_path = Path(sys.argv[1])
    if not md_path.is_file():
        print(f"error: {md_path} not found")
        return 1
    out_path = Path(sys.argv[2]) if len(sys.argv) > 2 \
        else md_path.with_suffix(".docx")
    convert(md_path, out_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
